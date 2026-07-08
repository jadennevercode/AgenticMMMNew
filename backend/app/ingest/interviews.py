"""Loader for stakeholder interview transcripts (.docx).

11 transcripts under ``纪要/``. Filenames encode interview layer & role, e.g.
``...Interview_Layer1_GM_20251119.docx`` -> layer="Layer1", role="GM".
Body text lives in the document paragraphs; a small metadata table (Date/Theme/
Attendee) is appended for completeness.
"""
from __future__ import annotations

import re
from pathlib import Path

import docx

from ._paths import INTERVIEW_DIR, ref

_LAYER_RE = re.compile(r"(Layer\s*\d)", re.IGNORECASE)
# Known role tokens that appear after the layer in filenames.
_ROLE_TOKENS = [
    "GM", "Finance", "Mkt", "Sales", "Media", "Activation",
    "RTM", "EC", "O2O", "SIA",
]


def _parse_layer_role(filename: str) -> tuple[str, str]:
    layer_match = _LAYER_RE.search(filename)
    layer = layer_match.group(1).replace(" ", "") if layer_match else ""
    role = ""
    for token in _ROLE_TOKENS:
        if re.search(rf"[_ ]{re.escape(token)}[_ .]", filename, re.IGNORECASE):
            role = token
            break
    return layer, role


def _docx_text(path: Path) -> str:
    doc = docx.Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_interviews() -> list[dict]:
    """Return one record per interview docx.

    Each record: {filename, layer, role, text} where ``text`` is the full
    transcript (paragraphs + metadata table). Returns records sorted by filename;
    unreadable files are skipped silently from the list but reported via the
    ``error`` key so callers can surface them.
    """
    out: list[dict] = []
    for path in sorted(ref(INTERVIEW_DIR).glob("*.docx")):
        if path.name.startswith("~$"):
            continue
        layer, role = _parse_layer_role(path.name)
        try:
            text = _docx_text(path)
            out.append({
                "filename": path.name,
                "layer": layer,
                "role": role,
                "text": text,
            })
        except Exception as exc:  # noqa: BLE001 - report, do not crash ingestion
            out.append({
                "filename": path.name,
                "layer": layer,
                "role": role,
                "text": "",
                "error": str(exc),
            })
    return out

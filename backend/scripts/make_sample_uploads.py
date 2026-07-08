"""Generate realistic sample upload files for manual UI testing of the upload gates.

S1 gates get document samples; the S2 data gate (2.0a) gets one workbook **per L3
slot**, derived from the real factor tree so each slot *validates* (sheet == L4,
headers == the L4's indicators) and the per-L3 coverage gate can actually clear.
Each L4 sheet carries the long-table granularity columns + its indicators, so the
files also parse into the per-project long table (the 2.21 schema binding).

    sample-uploads/
      project_background/   -> gate 1.0a
      industry_reference/   -> gate 1.1a
      interview_minutes/    -> gate 1.4a
      data/                 -> gate 2.0a  (one <L3>.xlsx per slot; upload each to its L3)

Run from the backend venv:
    PYTHONPATH=. .venv/bin/python scripts/make_sample_uploads.py
"""
from __future__ import annotations

import math
import random
import re
from pathlib import Path

import openpyxl
from docx import Document

from app import ingest

random.seed(20240612)

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "sample-uploads"

# The data-request EXPORT template header (Time + model-scope dims), so the sample
# matches what the client actually fills and what the binding parser reads.
TEMPLATE_DIMS = ["Time (Month)", "Product", "Channel", "Platform & Region"]
REGIONS = ["华东", "华南", "华北"]
CHANNEL_TYPES = ["MT", "TT", "EC", "O2O", "AFH"]
PERIODS = [f"{y}-{m:02d}" for y in (2023, 2024) for m in range(1, 13)]  # 24 monthly points
_INVALID = re.compile(r'[\\/:*?"<>|\[\]]+')


def _safe_filename(s: str) -> str:
    return _INVALID.sub("_", s).strip() or "slot"


def _safe_sheet(s: str) -> str:
    return _INVALID.sub("_", s).strip()[:31] or "sheet"


def _season(period_idx: int) -> float:
    """Shared annual seasonal multiplier (summer peak) applied to BOTH drivers and
    the KPI, so the response co-moves with the X drivers and OLS finds real signal."""
    return 1.0 + 0.32 * math.sin(2 * math.pi * (period_idx % 12) / 12.0)


def _val(indicator: str, season: float) -> float:
    """A plausible value scaled by the indicator kind, modulated by seasonality."""
    name = indicator.lower()
    if any(k in indicator for k in ("率", "占比", "%")) or "rate" in name:
        return round(min(1.0, max(0.0, 0.5 * season * random.uniform(0.85, 1.15))), 3)
    if any(k in indicator for k in ("花费", "spend", "金额", "RMB", "费用")):
        return round(800 * season * random.uniform(0.8, 1.2), 1)
    return round(20000 * season * random.uniform(0.85, 1.15), 1)


# Region / channel level effects for a believable KPI cross-section.
_REGION_EFF = {"华东": 1.25, "华南": 1.0, "华北": 0.8}
_CHANNEL_EFF = {"MT": 1.3, "TT": 1.0, "EC": 0.9, "O2O": 0.7, "AFH": 0.6}


def emit_data_slots() -> int:
    """One workbook per L3 (sheet==L4) + a KPI workbook carrying the response Y
    (本品月度销量), all in the export-template format. Drivers and Y share the same
    seasonal signal so the downstream OLS fit is real, not noise."""
    tree = ingest.load_factor_tree()["tree"]
    by_l3: dict[str, dict[str, list[str]]] = {}
    for _l1, l2s in tree.items():
        for _l2, l3s in l2s.items():
            for l3, l4s in l3s.items():
                bucket = by_l3.setdefault(l3 or "—", {})
                for l4, recs in l4s.items():
                    inds = [r["indicator"] for r in recs if r.get("indicator")]
                    bucket.setdefault(l4 or l3, [])
                    for ind in inds:
                        if ind not in bucket[l4 or l3]:
                            bucket[l4 or l3].append(ind)
    (OUT / "data").mkdir(parents=True, exist_ok=True)
    count = 0
    for l3, l4s in by_l3.items():
        wb = openpyxl.Workbook()
        wb.remove(wb.active)
        for l4, inds in l4s.items():
            ws = wb.create_sheet(_safe_sheet(l4))
            ws.append(TEMPLATE_DIMS + (inds or ["value"]))  # Time · Product · Channel · Region · indicators
            for rg in REGIONS:
                for ct in CHANNEL_TYPES:
                    for pi, period in enumerate(PERIODS):
                        s = _season(pi)
                        row = [period, "脉动 Mizone", ct, rg]
                        row += [_val(i, s) for i in (inds or ["value"])]
                        ws.append(row)
        path = OUT / "data" / f"{_safe_filename(l3)}.xlsx"
        wb.save(path)
        count += 1

    # KPI workbook — the dependent variable Y (本品月度销量), correlated with drivers
    # via the shared seasonal signal + region/channel effects + noise.
    wb = openpyxl.Workbook(); wb.remove(wb.active)
    ws = wb.create_sheet("本品销量")
    ws.append(TEMPLATE_DIMS + ["本品月度销量"])
    for rg in REGIONS:
        for ct in CHANNEL_TYPES:
            for pi, period in enumerate(PERIODS):
                s = _season(pi)
                base = 100000 * _REGION_EFF.get(rg, 1.0) * _CHANNEL_EFF.get(ct, 1.0)
                y = base * s * random.uniform(0.92, 1.08)  # noise small → fit is real
                ws.append([period, "脉动 Mizone", ct, rg, round(y, 1)])
    wb.save(OUT / "data" / "KPI_本品销量.xlsx")
    count += 1
    print(f"  + data/: {count} workbooks (export-template format) incl. KPI 本品月度销量 (Y)")
    return count


def _docx(path: Path, title: str, paras: list[str]) -> None:
    doc = Document()
    doc.add_heading(title, level=0)
    for p in paras:
        if p.startswith("# "):
            doc.add_heading(p[2:], level=1)
        elif p.startswith("- "):
            doc.add_paragraph(p[2:], style="List Bullet")
        else:
            doc.add_paragraph(p)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    print(f"  + {path.relative_to(ROOT)}")


def emit_s1_docs() -> None:
    _docx(OUT / "project_background" / "Mizone_MMM_SOW_signed.docx",
          "Statement of Work — Mizone MMM (Danone China)",
          ["# Engagement",
           "Brand: Mizone (脉动). Category: Functional / vitamin beverage. Market: China.",
           "Objective: Build a Marketing Mix Model to quantify media & trade ROI and "
           "recommend an FY25 budget reallocation.",
           "# Scope",
           "- Modelling window: weekly/monthly, ~104 weeks ending 2024-08.",
           "- Geography: 5 Nielsen regions. Channels: MT / TT / AFH / EC / O2O.",
           "- KPI: Nielsen weekly offtake volume.",
           "# Time granularity", "Monthly. National macro covariates included."])
    _docx(OUT / "project_background" / "Kickoff_brief.docx",
          "Kickoff Brief — Mizone MMM",
          ["# Business question",
           "Marketing spend grew 18% YoY but offtake was flat. Where is the money working?",
           "# Hypotheses",
           "- TV saturating in Tier-1 cities.", "- Douyin/KOL rising but unclear payback.",
           "- Summer temperature and distribution drive the seasonal swing."])
    _docx(OUT / "industry_reference" / "202408_Competitor_Benchmark.docx",
          "Competitor Benchmark — Functional Beverage (2024-08)",
          ["# Category", "Functional/vitamin water grew ~6% YoY; Mizone #1 in the hydration-vitamin sub-segment.",
           "# Key competitors",
           "- Gatorade: sports-hydration, heavy TV.", "- Pocari Sweat: electrolyte, summer flighting.",
           "- Genki Forest 外星人: Douyin-led, aggressive EC promo.",
           "# Pricing", "Shelf price 4.5–5.2 RMB/600ml; promo depth 15–25% during 618 / 11.11."])
    for name, role, notes in [
        ("Layer1_GM", "GM (Layer 1)",
         ["- 'We over-index on TV in summer; the last 20% of GRPs are likely wasted.'",
          "- Growth priority is EC and O2O; trade promo ROI is a concern.",
          "- Distribution quality and promo are must-have factors; KOL payback is the open question."]),
        ("Layer3_Media", "Media Team (Layer 3)",
         ["- TV bought on cost-per-GRP; OLV/Social for reach; Search always-on.",
          "- KOL flighted in bursts (~3 wks on, dark between).",
          "- Adstock longer for TV/OOH, short for Search/Social; flag KOL data gaps."]),
        ("Layer3_SIA", "Sales & In-store Activation (Layer 3)",
         ["- Feature & display drive most promo lift; distribution gains stick after summer.",
          "- 618 and 11.11 are the two mega EC events — separate them in the model.",
          "- Need promo flag, feature/display %, weighted distribution, mega-promo dummy."]),
    ]:
        _docx(OUT / "interview_minutes" / f"{name}.docx", f"Interview Minutes — {role}",
              ["# Notes", *notes])


def main() -> None:
    print(f"Writing sample uploads to {OUT.relative_to(ROOT)}/ …")
    print("data/ (gate 2.0a · per-L3 slots from the factor tree):")
    try:
        emit_data_slots()
    except Exception as exc:  # noqa: BLE001
        print(f"  ! factor tree unavailable ({exc}); skipping data slots")
    print("S1 document gates (1.0a / 1.1a / 1.4a):")
    emit_s1_docs()
    print("done.")


if __name__ == "__main__":
    main()
